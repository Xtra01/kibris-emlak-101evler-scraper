import os
from datetime import datetime
from typing import List, Optional

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from scraper.report import fetch_exchange_rates, add_computed_columns, DONUM_M2
from scraper import orchard_analysis as oa
from scraper import search


REPORTS_DIR = 'reports'


def _ensure_reports():
    os.makedirs(REPORTS_DIR, exist_ok=True)


def _fmt_try(x: Optional[float]) -> str:
    if pd.isna(x) or x is None:
        return '-'
    try:
        return f"₺{float(x):,.0f}".replace(',', '.')
    except Exception:
        return str(x)


def _stats_gt0(series: pd.Series) -> Optional[dict]:
    s = pd.to_numeric(series, errors='coerce')
    s = s[(~s.isna()) & (s > 0)]
    if s.empty:
        return None
    return {
        'count': int(s.count()),
        'min': float(s.min()),
        'p25': float(s.quantile(0.25)),
        'median': float(s.median()),
        'p75': float(s.quantile(0.75)),
        'max': float(s.max()),
        'mean': float(s.mean()),
    }


def _add_heading(doc: Document, text: str, level: int = 0):
    h = doc.add_heading(text, level=level)
    return h


def _add_paragraph(doc: Document, text: str, bold: bool = False, align_center: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_stats_table(doc: Document, title: str, st: Optional[dict]):
    _add_paragraph(doc, title, bold=True)
    if not st:
        _add_paragraph(doc, 'Veri bulunamadı.')
        return
    table = doc.add_table(rows=2, cols=7)
    hdr = table.rows[0].cells
    hdr[0].text = 'Kayıt'
    hdr[1].text = 'Min'
    hdr[2].text = 'P25'
    hdr[3].text = 'Medyan'
    hdr[4].text = 'P75'
    hdr[5].text = 'Max'
    hdr[6].text = 'Ort.'
    row = table.rows[1].cells
    row[0].text = str(st['count'])
    row[1].text = _fmt_try(st['min'])
    row[2].text = _fmt_try(st['p25'])
    row[3].text = _fmt_try(st['median'])
    row[4].text = _fmt_try(st['p75'])
    row[5].text = _fmt_try(st['max'])
    row[6].text = _fmt_try(st['mean'])


def _comps_table(doc: Document, df: pd.DataFrame, title: str, max_rows: int = 10):
    cols = ['property_id', 'district', 'area_m2', 'price_try', 'price_per_donum_try', 'url', 'title']
    available = [c for c in cols if c in df.columns]
    if not available:
        return
    _add_paragraph(doc, title, bold=True)
    view = df.copy()
    if 'area_m2' in view.columns:
        view['donum'] = (pd.to_numeric(view['area_m2'], errors='coerce') / DONUM_M2).round(2)
    view = view.head(max_rows)
    # create table
    table = doc.add_table(rows=len(view) + 1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = 'İlçe/Mahalle'
    hdr[1].text = 'Alan (dönüm)'
    hdr[2].text = 'Fiyat (₺)'
    hdr[3].text = '₺/dönüm'
    hdr[4].text = 'Başlık'
    hdr[5].text = 'Link'
    for i, (_, r) in enumerate(view.iterrows(), start=1):
        d = r.get('district', '')
        don = r.get('donum', '')
        pr = _fmt_try(r.get('price_try'))
        ppd = _fmt_try(r.get('price_per_donum_try'))
        t = str(r.get('title', ''))[:80]
        u = str(r.get('url', ''))
        row = table.rows[i].cells
        row[0].text = str(d) if pd.notna(d) else ''
        row[1].text = str(don)
        row[2].text = pr
        row[3].text = ppd
        row[4].text = t
        row[5].text = u


def build_agent_report(output_path: Optional[str] = None):
    _ensure_reports()
    # Load and compute
    df = pd.read_csv('property_details.csv')
    rates = fetch_exchange_rates()
    df = add_computed_columns(df, rates)

    # Baseline Güzelyurt arsa ≥1 dönüm
    base = search.advanced_search(city='guzelyurt', property_type='arsa', listing_type='Sale', min_donum=1, limit=5000)

    # Text for orchard/water detection (same logic as orchard_analysis)
    text = (base['title'].fillna('') + ' ' + base['description'].fillna('')).map(oa._norm_text)
    orchard_terms = [
        'portakal','narenciye','mandalina','limon','meyve','bahce','bahcesi','bahcedir','agac','zeytin','ceviz','incir'
    ]
    water_terms = [
        'tatli su','tc su','turkiye su','turkiyeden su','anavatandan su','sulama','sulama suyu','su hatti','su baglanti','su saati',
        'sebeke suyu','kuyu','artezyen','artezyen kuyu','damla','damlama','hidrofor','su deposu','depo','tarimsal su','ciftci suyu','belediye su'
    ]
    mask_orchard = oa.any_token_in(text, orchard_terms)
    mask_water = oa.any_token_in(text, water_terms)

    orchard_df = base[mask_orchard].copy()
    orchard_water_df = base[mask_orchard & mask_water].copy()

    # Core focus: Güzelyurt Merkez / Piyalepaşa
    dist_series = base.get('district', pd.Series('', index=base.index)).fillna('').map(oa._norm_text)
    mask_core = dist_series.str.contains('guzelyurt') & (dist_series.str.contains('merkez') | dist_series.str.contains('piyalepasa'))
    orchard_core_df = orchard_df[mask_core.loc[orchard_df.index]].copy()

    # 10+ dönüm scenario in Güzelyurt (for larger parcel comps)
    large = search.advanced_search(city='guzelyurt', property_type='arsa', listing_type='Sale', min_donum=10, limit=5000)
    orchard_large = large[oa.any_token_in((large['title'].fillna('') + ' ' + large['description'].fillna('')).map(oa._norm_text), orchard_terms)].copy()

    # Stats (exclude zeros to avoid data-entry artifacts)
    st_base = _stats_gt0(base['price_per_donum_try'])
    st_orch = _stats_gt0(orchard_df['price_per_donum_try'])
    st_orch_core = _stats_gt0(orchard_core_df['price_per_donum_try'])
    st_large = _stats_gt0(large['price_per_donum_try'])
    st_orch_large = _stats_gt0(orchard_large['price_per_donum_try'])

    # Recommended pricing band logic
    # Use orchard median as anchor; apply micro-location premium if core stats exist, but note sample size.
    median_base = st_base['median'] if st_base else None
    median_orch = st_orch['median'] if st_orch else median_base
    median_core = st_orch_core['median'] if st_orch_core else None

    # Default band around orchard median ± ~10-15%, bounded by base p25/p75 when available
    if median_orch:
        low = median_orch * 0.95
        high = median_orch * 1.15
        # If core shows premium and is not wildly off, nudge band upward slightly
        if median_core and median_core > median_orch * 1.5:
            high = max(high, median_orch * 1.2)
    else:
        low, high = None, None

    # Build DOCX
    doc = Document()
    _add_paragraph(doc, 'Güzelyurt – Piyalepaşa 10 Dönüm Narenciye Arazisi', bold=True, align_center=True)
    _add_paragraph(doc, 'Piyasa Değer Analizi ve Fiyatlandırma Önerisi', align_center=True)
    _add_paragraph(doc, datetime.now().strftime('%d %B %Y'), align_center=True)

    doc.add_paragraph()
    _add_heading(doc, '1) Yönetici Özeti', level=1)
    reco_lines: List[str] = []
    if low and high:
        reco_lines.append(f"Önerilen listeleme bandı: {_fmt_try(low)} – {_fmt_try(high)} / dönüm")
        reco_lines.append(f"10 dönüm toplam hedef değer: {_fmt_try(low*10)} – {_fmt_try(high*10)}")
    else:
        reco_lines.append("Öneri: Narenciye medyanına dayalı birim fiyat bandı, detaylar aşağıda.")
    reco_lines.append("Temellendirme: Güzelyurt satılık arsa ≥1 dönüm medyanı, narenciye eşleşmeleri medyanı, çekirdek (Merkez/Piyalepaşa) premiumu ve 10+ dönüm segmenti dikkate alınmıştır.")
    _add_paragraph(doc, '\n'.join(['• ' + x for x in reco_lines]))

    doc.add_paragraph()
    _add_heading(doc, '2) Kanıt ve Piyasa Özeti', level=1)
    _add_paragraph(doc, f"Genel Güzelyurt arsa (≥1 dönüm): {st_base['count'] if st_base else 0} ilan; medyan {_fmt_try(st_base['median']) if st_base else '-'} / dönüm.")
    _add_paragraph(doc, f"Narenciye/bahçe/ağaç eşleşmeleri: {st_orch['count'] if st_orch else 0} ilan; medyan {_fmt_try(st_orch['median']) if st_orch else '-'} / dönüm.")
    _add_paragraph(doc, f"Çekirdek (Güzelyurt Merkez/Piyalepaşa) narenciye: {st_orch_core['count'] if st_orch_core else 0} ilan; medyan {_fmt_try(st_orch_core['median']) if st_orch_core else '-'} / dönüm (küçük örnek uyarısı).")
    _add_paragraph(doc, f"Geniş parseller (≥10 dönüm) Güzelyurt: {st_large['count'] if st_large else 0} ilan; medyan {_fmt_try(st_large['median']) if st_large else '-'} / dönüm.")
    _add_paragraph(doc, f"Narenciye (≥10 dönüm) eşleşmeleri: {st_orch_large['count'] if st_orch_large else 0} ilan; medyan {_fmt_try(st_orch_large['median']) if st_orch_large else '-'} / dönüm.")

    # Detailed stats tables
    doc.add_paragraph()
    _add_stats_table(doc, 'Genel Güzelyurt Arsa (≥1 dönüm) — ₺/dönüm', st_base)
    doc.add_paragraph()
    _add_stats_table(doc, 'Narenciye Eşleşmeleri — ₺/dönüm', st_orch)
    doc.add_paragraph()
    _add_stats_table(doc, 'Çekirdek (Merkez/Piyalepaşa) Narenciye — ₺/dönüm', st_orch_core)
    doc.add_paragraph()
    _add_stats_table(doc, '≥10 Dönüm — ₺/dönüm', st_large)

    # Comparables: prioritize orchard >=8 dönüm; fallback to base >=8 dönüm sorted by price_per_donum_try
    comps = orchard_large.copy()
    if comps.empty:
        comps = large.copy()
    comps = comps[pd.to_numeric(comps.get('price_per_donum_try'), errors='coerce') > 0]
    comps = comps.sort_values(by='price_per_donum_try', ascending=True, na_position='last')
    if not comps.empty:
        doc.add_paragraph()
        _comps_table(doc, comps, 'Karşılaştırılabilir İlanlar (≥10 dönüm, ₺/dönüm artan)', max_rows=12)

    # Positioning & negotiation
    doc.add_paragraph()
    _add_heading(doc, '3) Konumlandırma ve Pazarlık Stratejisi', level=1)
    bullets = [
        "Fiyat çapasını narenciye medyanı üzerine kurup, çekirdek bölge primini 'yerel nadirlik' argümanıyla destekleyin.",
        "Su altyapısı metinde nadiren yazıldığı için ilan linkleri ve telefon teyidiyle ispatlayın; bu, alıcı gözünde risk indirir (fiyatı destekler).",
        "Ağaçların kuruluğu/rehabilitasyon ihtiyacı için net maliyet kalemleri (söküm, yeni dikim, damlama) çıkarın; pazarlıkta bu kalemleri fiyata yansıtın.",
        "Listeleme fiyatını önerilen bandın üst sınırına yakın başlatıp, deneme gösterimleriyle 2-3 hafta içinde pazarın tepkisine göre optimize edin.",
        "Alıcıya 'çekirdek mikro-lokasyon' ile 'geniş parsel arz kıtlığı'nı aynı anda göstererek fiyat aralığımızın rasyonel olduğunu kanıtlayın."
    ]
    _add_paragraph(doc, '\n'.join(['• ' + b for b in bullets]))

    # Save
    if not output_path:
        output_path = os.path.join(REPORTS_DIR, f"Orchard_Valuation_Piyalepasa_{datetime.now().strftime('%Y%m%d')}.docx")
    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    out = build_agent_report()
    print(out)
